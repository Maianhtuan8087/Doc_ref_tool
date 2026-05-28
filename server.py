from __future__ import annotations

import cgi
import copy
import re
import sys
import tempfile
import unicodedata
from dataclasses import dataclass
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import quote

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


APP_DIR = Path(__file__).resolve().parent
INDEX_PATH = APP_DIR / "index.html"
REF_HEADING_RE = re.compile(r"^\s*tài\s+liệu\s+tham\s+khảo\b", re.IGNORECASE)
REF_ITEM_RE = re.compile(r"^\s*(?:\[(\d+)\]|(\d+)[\.\)]|\((\d+)\))\s*(.*)$")
CITATION_RE = re.compile(r"\[((?:\d+\s*(?:[-–]\s*\d+)?)(?:\s*,\s*\d+\s*(?:[-–]\s*\d+)?)*)\]")


@dataclass
class SourceDoc:
    path: Path
    doc: Document
    body_elements: list
    refs: list[tuple[int, str]]
    citation_map: dict[int, int]


@dataclass
class RefRecord:
    text: str
    language: str
    first_seen: int


def is_ref_heading(text: str) -> bool:
    return bool(REF_HEADING_RE.match(text.strip()))


def clean_reference_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    return re.sub(r"\s+\[(?:\d+(?:\s*[-–,]\s*\d+)*)\]\s*$", "", text)


def strip_ref_number(text: str, fallback_number: int) -> tuple[int, str, bool]:
    match = REF_ITEM_RE.match(text)
    if not match:
        return fallback_number, text.strip(), False

    number = int(next(group for group in match.groups()[:3] if group))
    return number, match.group(4).strip(), True


def parse_references(lines: list[str]) -> tuple[list[tuple[int, str]], int]:
    refs: list[tuple[int, str]] = []
    next_number = 1
    unnumbered_count = 0

    for line in lines:
        text = clean_reference_text(line)
        if not text:
            continue

        number, ref_text, explicit = strip_ref_number(text, next_number)
        if not explicit:
            unnumbered_count += 1

        ref_text = clean_reference_text(ref_text)
        if ref_text:
            refs.append((number, ref_text))
            next_number = max(next_number, number + 1)

    return refs, unnumbered_count


def split_doc(path: Path) -> tuple[Document, list, list[tuple[int, str]], int]:
    doc = Document(path)
    body_elements = []
    ref_lines: list[str] = []
    in_refs = False

    for element in doc.element.body.iterchildren():
        if element.tag == qn("w:sectPr"):
            continue

        if element.tag == qn("w:p"):
            paragraph = Paragraph(element, doc)
            text = paragraph.text.strip()
            if not in_refs and is_ref_heading(text):
                in_refs = True
                continue
            if in_refs:
                ref_lines.append(text)
            else:
                body_elements.append(element)
        elif in_refs:
            continue
        else:
            body_elements.append(element)

    refs, unnumbered_count = parse_references(ref_lines)
    return doc, body_elements, refs, unnumbered_count


def normalize_ref(text: str) -> str:
    text = re.sub(r"[\W_]+", " ", text.casefold(), flags=re.UNICODE)
    return re.sub(r"\s+", " ", text).strip()


def strip_accents(text: str) -> str:
    normalized = unicodedata.normalize("NFD", text)
    return "".join(char for char in normalized if unicodedata.category(char) != "Mn")


def classify_reference_language(text: str) -> str:
    vietnamese_chars = "ăâđêôơưáàảãạấầẩẫậắằẳẵặéèẻẽẹếềểễệíìỉĩịóòỏõọốồổỗộớờởỡợúùủũụứừửữựýỳỷỹỵ"
    lowered = text.casefold()
    if any(char in lowered for char in vietnamese_chars):
        return "vi"

    plain = re.sub(r"[\W_]+", " ", strip_accents(lowered))
    vietnamese_terms = [
        "bo y te",
        "quyet dinh",
        "huong dan",
        "ha noi",
        "benh",
        "dieu tri",
        "tang huyet ap",
        "tram y te",
    ]
    if any(term in plain for term in vietnamese_terms):
        return "vi"

    return "en"


def parse_citation_numbers(citation_body: str) -> list[int]:
    numbers: list[int] = []
    for part in re.split(r"\s*,\s*", citation_body.strip()):
        range_match = re.match(r"^(\d+)\s*[-–]\s*(\d+)$", part)
        if range_match:
            start, end = int(range_match.group(1)), int(range_match.group(2))
            step = 1 if end >= start else -1
            numbers.extend(range(start, end + step, step))
        elif part.isdigit():
            numbers.append(int(part))
    return numbers


def compact_numbers(numbers: list[int]) -> str:
    unique: list[int] = []
    for number in numbers:
        if number not in unique:
            unique.append(number)

    chunks: list[str] = []
    i = 0
    while i < len(unique):
        start = unique[i]
        end = start
        while i + 1 < len(unique) and unique[i + 1] == end + 1:
            i += 1
            end = unique[i]
        if end - start >= 2:
            chunks.append(f"{start}-{end}")
        else:
            chunks.extend(str(number) for number in range(start, end + 1))
        i += 1
    return ", ".join(chunks)


def iter_paragraphs(element, doc: Document):
    for paragraph_element in element.iter(qn("w:p")):
        yield Paragraph(paragraph_element, doc)


def replace_citations_in_paragraph(paragraph: Paragraph, citation_map: dict[int, int], warnings: set[str], source_name: str) -> None:
    def replace(match: re.Match[str]) -> str:
        old_numbers = parse_citation_numbers(match.group(1))
        if not old_numbers:
            return match.group(0)

        missing = [number for number in old_numbers if number not in citation_map]
        if missing:
            warnings.add(f"{source_name}: citation {match.group(0)} không có mục tham khảo tương ứng")
            return match.group(0)

        return f"[{compact_numbers([citation_map[number] for number in old_numbers])}]"

    for run in paragraph.runs:
        if "[" in run.text and "]" in run.text:
            run.text = CITATION_RE.sub(replace, run.text)


def clear_document_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def append_body_element(doc: Document, element) -> None:
    body = doc.element.body
    insert_at = len(body) - 1 if body.sectPr is not None else len(body)
    body.insert(insert_at, copy.deepcopy(element))


def process_docx(paths: list[Path], output_path: Path) -> str:
    warnings: set[str] = set()
    sources: list[SourceDoc] = []
    ref_records: list[RefRecord] = []
    ref_index_maps: list[dict[int, int]] = []
    report_lines: list[str] = []

    for path in paths:
        doc, body_elements, refs, unnumbered_count = split_doc(path)
        ref_index_map: dict[int, int] = {}

        for old_number, ref_text in refs:
            ref_index = len(ref_records)
            ref_records.append(
                RefRecord(
                    text=ref_text,
                    language=classify_reference_language(ref_text),
                    first_seen=ref_index,
                )
            )
            ref_index_map[old_number] = ref_index

        sources.append(SourceDoc(path, doc, body_elements, refs, {}))
        ref_index_maps.append(ref_index_map)

        note = f"{path.name}: {len(refs)} tài liệu tham khảo"
        if unnumbered_count:
            note += f", {unnumbered_count} mục không có số gốc được đánh số theo thứ tự xuất hiện"
        report_lines.append(note)

    if not ref_records:
        raise ValueError("Không tìm thấy phần 'Tài liệu tham khảo' trong file DOCX.")

    vietnamese_refs = [record for record in ref_records if record.language == "vi"]
    english_refs = [record for record in ref_records if record.language == "en"]
    final_refs = vietnamese_refs + english_refs
    final_number_by_first_seen = {record.first_seen: index for index, record in enumerate(final_refs, start=1)}

    for source, ref_index_map in zip(sources, ref_index_maps):
        source.citation_map = {
            old_number: final_number_by_first_seen[ref_index]
            for old_number, ref_index in ref_index_map.items()
        }
        for element in source.body_elements:
            for paragraph in iter_paragraphs(element, source.doc):
                replace_citations_in_paragraph(paragraph, source.citation_map, warnings, source.path.name)

    output_doc = Document(paths[0])
    clear_document_body(output_doc)

    for source_index, source in enumerate(sources):
        for element in source.body_elements:
            append_body_element(output_doc, element)
        if source_index != len(sources) - 1:
            output_doc.add_paragraph()

    output_doc.add_paragraph()
    output_doc.add_paragraph("Tài liệu tham khảo")
    output_doc.add_paragraph("Tài liệu tiếng Việt")
    for index, ref in enumerate(vietnamese_refs, start=1):
        output_doc.add_paragraph(f"{index}. {ref.text}")

    output_doc.add_paragraph("Tài liệu tiếng Anh")
    for index, ref in enumerate(english_refs, start=len(vietnamese_refs) + 1):
        output_doc.add_paragraph(f"{index}. {ref.text}")

    output_doc.save(output_path)

    report_lines.append("")
    report_lines.append(f"Tổng số tài liệu tham khảo sau xử lý: {len(final_refs)}")
    report_lines.append(f"Tài liệu tiếng Việt: {len(vietnamese_refs)}")
    report_lines.append(f"Tài liệu tiếng Anh: {len(english_refs)}")
    if warnings:
        report_lines.append("")
        report_lines.append("Cảnh báo:")
        report_lines.extend(f"- {warning}" for warning in sorted(warnings))
    return "\n".join(report_lines)


class Handler(BaseHTTPRequestHandler):
    server_version = "DocxRefTool/1.0"

    def do_GET(self) -> None:
        if self.path not in {"/", "/index.html"}:
            self.send_error(404)
            return

        data = INDEX_PATH.read_bytes()
        self.send_response(200)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def do_POST(self) -> None:
        if self.path != "/api/process":
            self.send_error(404)
            return

        try:
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )
            mode = form.getfirst("mode", "merge")
            file_fields = form["files"] if "files" in form else []
            if not isinstance(file_fields, list):
                file_fields = [file_fields]
            file_fields = [field for field in file_fields if getattr(field, "filename", None)]

            if not file_fields:
                self.respond_text(400, "Chưa có file DOCX nào được upload.")
                return

            with tempfile.TemporaryDirectory(prefix="docx_ref_tool_") as tmp:
                tmp_dir = Path(tmp)
                input_paths: list[Path] = []
                for index, field in enumerate(file_fields, start=1):
                    name = Path(field.filename).name
                    if not name.lower().endswith(".docx"):
                        self.respond_text(400, f"File không phải DOCX: {name}")
                        return
                    path = tmp_dir / f"{index:03d}_{name}"
                    path.write_bytes(field.file.read())
                    input_paths.append(path)

                output_name = "tach_tai_lieu_viet_anh.docx" if mode == "split" else "ghep_docx_tai_lieu_tham_khao.docx"
                output_path = tmp_dir / output_name
                report = process_docx(input_paths, output_path)
                data = output_path.read_bytes()

            self.send_response(200)
            self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header("Content-Disposition", f'attachment; filename="{output_name}"')
            self.send_header("Content-Length", str(len(data)))
            self.send_header("X-Process-Report", quote(report))
            self.end_headers()
            self.wfile.write(data)
        except Exception as exc:
            self.respond_text(500, f"Lỗi xử lý: {exc}")

    def respond_text(self, status: int, text: str) -> None:
        data = text.encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "text/plain; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)


def main() -> int:
    sys.stdout.reconfigure(encoding="utf-8")
    port = 8765
    server = ThreadingHTTPServer(("127.0.0.1", port), Handler)
    print(f"DOCX Reference Tool: http://127.0.0.1:{port}")
    print("Nhấn Ctrl+C để dừng server.")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nĐã dừng server.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
